import streamlit as st
import datetime
import google.generativeai as genai
import time
import google.api_core.exceptions
import re
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import base64

# HAPUS googletrans dan gunakan alternatif
try:
    from deep_translator import GoogleTranslator
    TRANSLATOR_AVAILABLE = True
except ImportError:
    TRANSLATOR_AVAILABLE = False

# Data default Kevin
DEFAULT_DATA = {
    "nama": "Kevin Jonathan Sinaga", 
    "telepon": "08999714652",
    "email": "kevinjonathans01@email.com",
    "alamat": "Jl Berdikari, Medan",
    "linkedin": "https://www.linkedin.com/in/kevin-jonathan-a06a99292/",
    "portofolio": "https://kevins012.github.io/",
    "jurusan": "Teknik Elektro",
    "universitas": "Universitas Sumatera Utara"
}

# KEAHLIAN DEFAULT YANG BISA DIEDIT
DEFAULT_SKILLS = {
    "keahlian_teknis": [
        "Python Programming", 
        "Data Analysis", 
        "Web Development", 
        "SQL", 
        "Database Management",
        "Machine Learning Basics",
        "API Development"
    ],
    "keahlian_lunak": [
        "Problem Solving", 
        "Team Leadership", 
        "Communication", 
        "Fast Learner", 
        "Adaptability",
        "Project Management",
        "Critical Thinking"
    ]
}

# PENGALAMAN DEFAULT YANG BISA DIEDIT
DEFAULT_EXPERIENCES = {
    "pengalaman_akademik": [
    "Mengembangkan robot sepak bola beroda: membangun sistem komunikasi antar-robot, odometri, integrasi sensor kompas, dan deteksi objek menggunakan image processing.",
    "Membangun sistem AI untuk menerjemahkan bahasa isyarat berbasis dua arsitektur neural network (time series dan non-time series) sebagai skripsi.",
    "Merancang platform web untuk sistem deteksi emisi kendaraan: mengelola data kepemilikan dan menerima input real-time dari Arduino R4.",
    "Terlibat dalam proyek penelitian sistem trading karbon berbasis data emisi kendaraan."
],
    "pengalaman_organisasi": [
        "Ketua Tim Robotika Kampus - memimpin 15 anggota dalam pengembangan robot sepak bola beroda",
   
    ],
    "pencapaian_lain": [
        "Sertifikasi dari deepLearning.AI selapa pembelajaran di Coursera",
        "Sertifikat SCADA",
       
    ]
}

# KONTEN DEFAULT YANG SUDAH PROFESIONAL
DEFAULT_CONTENT = {
    "pembuka": "Saya mengajukan lamaran untuk posisi {posisi} di {perusahaan} sebagai lulusan {jurusan} dari {universitas}, dengan fondasi akademis yang solid dan tekad kuat untuk berkembang serta memberikan kontribusi nyata dalam lingkungan kerja yang dinamis dan berorientasi pada hasil",
    
    "pencapaian": "Selama masa studi, saya berhasil mengembangkan sistem analisis data untuk tugas akhir yang mampu memproses dataset 10,000+ record. Saya juga memimpin tim project dalam kompetisi coding nasional dimana kami meraih peringkat 3 terbaik. Pengalaman ini melatih kemampuan teknis sekaligus leadership dalam menyelesaikan masalah kompleks.",
    
    "alasan_perusahaan": "Saya sangat mengagumi reputasi {perusahaan} dalam hal inovasi teknologi dan komitmen terhadap pengembangan solusi digital yang berdampak. Visi perusahaan dalam mendorong transformasi digital sejalan dengan passion saya dalam menciptakan solusi teknologi yang meaningful.",
    
    "motivasi_penutup": "Saya adalah pribadi yang cepat beradaptasi dan memiliki semangat belajar tinggi. Dengan dasar keahlian teknis yang solid dan kemampuan problem-solving yang terasah, saya yakin dapat memberikan kontribusi signifikan bagi {perusahaan}. Saya sangat antusias untuk bergabung dan berkembang bersama tim yang profesional."
}

# PROMPT SUPER DETAIL PROFESIONAL DENGAN ETHICAL CONSTRAINTS
PROFESSIONAL_PROMPTS = {
    "analyze_match": """
ANALYSIS TASK: Comprehensive Professional Match Analysis between Cover Letter and Job Description

CONTEXT:
- Position: {posisi}
- Company: {perusahaan}
- Applicant: {nama}
- Applicant Skills: {applicant_skills}
- Applicant Experiences: {applicant_experiences}
- Cover Letter Length: {cl_length} characters
- Job Description Length: {jd_length} characters

ANALYSIS FRAMEWORK - WEIGHTED SCORING (100 points total):

1. TECHNICAL SKILLS ALIGNMENT (25 points)
   • Keyword Density & Relevance (10pts)
   • Technical Competency Match (8pts)
   • Tool/Technology Proficiency (7pts)

2. EXPERIENCE & ACHIEVEMENT RELEVANCE (25 points)
   • Quantifiable Achievements (10pts)
   • Project Relevance (8pts)
   • Leadership & Initiative (7pts)

3. CULTURAL & VALUES FIT (20 points)
   • Company Research Depth (8pts)
   • Values Alignment (6pts)
   • Team Collaboration Fit (6pts)

4. MOTIVATION & CAREER NARRATIVE (15 points)
   • Career Story Coherence (6pts)
   • Passion Demonstration (5pts)
   • Long-term Alignment (4pts)

5. PROFESSIONAL IMPACT & PERSUASION (15 points)
   • Value Proposition Clarity (6pts)
   • Call-to-Action Effectiveness (5pts)
   • Professional Tone (4pts)

JOB DESCRIPTION KEY ELEMENTS:
{job_description_analysis}

COVER LETTER CONTENT:
{cover_letter_content}

ETHICAL ANALYSIS CONSTRAINTS:
- DO NOT suggest adding skills the applicant doesn't have
- Focus on TRANSFERABLE skills from existing experiences
- Highlight LEARNING CAPACITY and ADAPTABILITY
- Suggest REALISTIC skill development paths

ANALYSIS INSTRUCTIONS:
1. Score each category 0-100% then convert to weighted points
2. Provide SPECIFIC evidence from both documents
3. Identify TRANSFERABLE skills from applicant's background
4. Highlight LEARNING POTENTIAL for missing requirements
5. Give ACTIONABLE improvement recommendations

REQUIRED OUTPUT FORMAT:
OVERALL MATCH SCORE: [X]/100

DETAILED BREAKDOWN:
🔧 TECHNICAL ALIGNMENT: [X]/25
   • Strengths: [Specific examples from applicant's actual skills]
   • Transferable Skills: [Skills that can be applied to this role]
   • Learning Opportunities: [Realistic areas for growth]

📈 EXPERIENCE RELEVANCE: [X]/25
   • Achievement Alignment: [Specific matches from actual experiences]
   • Project Relevance: [How existing projects relate to position]
   • Leadership Experience: [Actual leadership examples]

🏢 CULTURAL FIT: [X]/20
   • Company Research: [Evidence of research]
   • Values Connection: [Specific value alignments]
   • Team Collaboration: [Actual teamwork experiences]

🎯 MOTIVATION NARRATIVE: [X]/15
   • Story Strength: [Career narrative evaluation]
   • Passion Evidence: [Specific passion demonstrations]
   • Alignment Clarity: [Future goal alignment]

💼 PROFESSIONAL IMPACT: [X]/15
   • Value Proposition: [Clarity of value offered based on actual skills]
   • Persuasion Effectiveness: [Call-to-action strength]
   • Professionalism: [Tone and language evaluation]

TRANSFERABLE SKILLS HIGHLIGHT:
- [Skill 1 from applicant's background that applies]
- [Skill 2 from applicant's background that applies]
- [Skill 3 from applicant's background that applies]

REALISTIC IMPROVEMENTS (Priority Order):
1. [Highest priority improvement - specific and actionable based on actual capabilities]
2. [Second priority improvement - specific and actionable]
3. [Third priority improvement - specific and actionable]

EXPECTED INTERVIEW SCORE: [X]/10 - [Brief rationale based on actual qualifications]
""",

    "enhance_cover": """
ENHANCEMENT TASK: Professional Cover Letter Optimization for Maximum Impact

CONTEXT:
- Target Position: {posisi}
- Target Company: {perusahaan}
- Applicant: {nama}
- Applicant Actual Skills: {applicant_skills}
- Applicant Actual Experiences: {applicant_experiences}
- Industry: Technology/IT
- Experience Level: Entry-Level Professional

ORIGINAL COVER LETTER:
{original_cover_letter}

JOB DESCRIPTION ANALYSIS:
{job_description_analysis}

MATCH ANALYSIS FINDINGS:
{match_analysis}

ETHICAL ENHANCEMENT CONSTRAINTS:
- DO NOT add skills or experiences the applicant doesn't actually have
- Focus on highlighting TRANSFERABLE skills from existing background
- Emphasize LEARNING CAPACITY and ADAPTABILITY
- Frame gaps as OPPORTUNITIES FOR GROWTH, not deficiencies
- Maintain HONESTY while maximizing professional appeal

ENHANCEMENT STRATEGY - PROFESSIONAL STANDARDS:

1. EXECUTIVE SUMMARY OPTIMIZATION (First Paragraph)
   • Immediate value proposition BASED ON ACTUAL SKILLS
   • Key achievement highlight FROM REAL EXPERIENCES
   • Direct position relevance WITH HONEST ASSESSMENT

2. ACHIEVEMENT QUANTIFICATION & IMPACT
   • Convert ACTUAL responsibilities to achievements
   • Add metrics and measurable results FROM REAL PROJECTS
   • Show business impact OF ACTUAL WORK

3. TRANSFERABLE SKILLS HIGHLIGHTING
   • Natural emphasis on EXISTING skills that apply
   • Show LEARNING POTENTIAL for required skills
   • Demonstrate ADAPTABILITY from past experiences

4. COMPANY-SPECIFIC CUSTOMIZATION
   • Demonstrate deep company research
   • Align with company mission/values
   • Show industry awareness

5. PROFESSIONAL NARRATIVE FLOW
   • Clear career story progression BASED ON FACTS
   • Logical argument building WITH INTEGRITY
   • Compelling conclusion THAT'S AUTHENTIC

ENHANCEMENT CRITERIA:

TRUTHFUL ACHIEVEMENTS:
- Quantify ONLY actual accomplishments
- Use specific numbers/percentages FROM REAL DATA
- Show actual impact of completed projects

TRANSFERABLE SKILLS INTEGRATION:
- Highlight 5-8 ACTUAL skills that match requirements
- Show how existing skills apply to new context
- Emphasize learning agility and adaptability

PROFESSIONAL FRAMING OF GAPS:
- Frame missing skills as learning opportunities
- Show enthusiasm for professional development
- Highlight foundational skills that enable quick learning

PROFESSIONAL TONE:
- Executive-level language
- Confident but not arrogant
- Enthusiastic but professional
- Honest but optimistic

STRUCTURAL OPTIMIZATION:
- Powerful opening hook BASED ON REAL STRENGTHS
- Logical paragraph transitions
- Memorable closing statement

OUTPUT REQUIREMENTS:
- Maintain original length ±15%
- Preserve all essential personal information
- Use professional business letter format
- Ensure ATS compatibility
- Create compelling narrative flow WITH INTEGRITY

FINAL ENHANCED COVER LETTER MUST:
1. Immediately capture recruiter attention WITH AUTHENTIC STRENGTHS
2. Demonstrate clear value proposition BASED ON ACTUAL SKILLS
3. Show quantifiable achievements FROM REAL EXPERIENCES
4. Exhibit deep company understanding
5. Create compelling call-to-action

IMPORTANT: Return ONLY the enhanced cover letter text without any explanations or additional text.
""",

    "improve_section": """
SECTION OPTIMIZATION TASK: Professional Content Enhancement WITH INTEGRITY

CONTEXT:
- Position: {posisi}
- Company: {perusahaan}
- Section Type: {section_type}
- Applicant Actual Skills: {applicant_skills}
- Applicant Actual Experiences: {applicant_experiences}
- Current Content Length: {content_length} characters

ORIGINAL CONTENT:
{original_content}

JOB DESCRIPTION KEY REQUIREMENTS:
{job_keywords}

ETHICAL CONSTRAINTS:
- DO NOT invent skills or experiences
- Focus on enhancing presentation of ACTUAL capabilities
- Highlight TRANSFERABLE skills from existing background
- Frame professionally without exaggeration

OPTIMIZATION GOALS for {section_type}:

FOR OPENING PARAGRAPH:
• Create immediate engagement WITH AUTHENTIC STRENGTHS
• Establish value proposition BASED ON REAL SKILLS
• Show enthusiasm and relevance HONESTLY

FOR ACHIEVEMENTS SECTION:
• Quantify ACTUAL accomplishments
• Show progressive responsibility FROM REAL EXPERIENCES
• Demonstrate impact and results OF ACTUAL WORK

FOR COMPANY ALIGNMENT:
• Show specific research
• Connect personal values to company
• Demonstrate cultural fit THROUGH ACTUAL EXPERIENCES

FOR CLOSING PARAGRAPH:
• Reinforce value proposition BASED ON FACTS
• Create urgency for interview
• Show professional confidence WITH INTEGRITY

ENHANCEMENT TECHNIQUES TO APPLY:

1. POWER WORD IMPLEMENTATION (FOR ACTUAL ACHIEVEMENTS):
   - Action verbs: Orchestrated, Spearheaded, Engineered (FOR REAL PROJECTS)
   - Achievement words: Accelerated, Optimized, Transformed (FOR ACTUAL RESULTS)
   - Leadership words: Mentored, Coordinated, Championed (FOR REAL EXPERIENCES)

2. METRICS INTEGRATION (FROM REAL DATA):
   - Percentage improvements FROM ACTUAL PROJECTS
   - Time/quantity metrics FROM REAL WORK
   - Scale indicators OF ACTUAL ACHIEVEMENTS

3. PROFESSIONAL LANGUAGE:
   - Industry-specific terminology
   - Business impact language
   - Executive-level vocabulary

4. TRANSFERABLE SKILLS STORYTELLING:
   - Show how existing skills apply to new role
   - Highlight learning capacity from past experiences
   - Demonstrate adaptability through real examples

OUTPUT REQUIREMENTS:
- Maintain section purpose and intent
- Enhance professional impact WITH INTEGRITY
- Integrate relevant JD keywords naturally
- Improve readability and flow
- Increase persuasive power HONESTLY

Return ONLY the enhanced section text without any explanations.
"""
}

def generate_with_gemini(prompt):
    """Generate konten dengan Gemini AI"""
    max_retries = 3
    delay = 5
    
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel('models/gemini-pro-latest')
            response = model.generate_content(prompt)
            return response.text.strip()
        except google.api_core.exceptions.ResourceExhausted as e:
            st.warning(f"Quota terlampaui (Percobaan {attempt + 1}/{max_retries}). Mencoba lagi dalam {delay} detik...")
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            st.error(f"Error Gemini: {str(e)}")
            return None
    
    st.error("Gagal mendapatkan respon dari AI setelah beberapa kali percobaan.")
    return None

def create_docx_file(cover_letter, perusahaan, posisi, nama):
    """Create DOCX file dengan formatting profesional"""
    doc = Document()
    
    # Judul dokumen
    title = doc.add_heading(f'Cover Letter - {nama}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informasi pribadi
    doc.add_paragraph(nama)
    doc.add_paragraph(DEFAULT_DATA['alamat'])
    doc.add_paragraph(f"{DEFAULT_DATA['telepon']} | {DEFAULT_DATA['email']} | {DEFAULT_DATA['linkedin']}")
    doc.add_paragraph(f"Portfolio: {DEFAULT_DATA['portofolio']}")
    
    doc.add_paragraph()  # Spasi
    
    # Tanggal
    date_paragraph = doc.add_paragraph(datetime.date.today().strftime("%d %B %Y"))
    
    doc.add_paragraph()  # Spasi
    
    # Penerima
    doc.add_paragraph("Kepada Yth:")
    doc.add_paragraph("Manajer Rekrutmen")
    doc.add_paragraph(perusahaan)
    
    doc.add_paragraph()  # Spasi
    
    # Subjek
    subject = doc.add_paragraph(f"Subjek: Lamaran Posisi {posisi}")
    
    doc.add_paragraph()  # Spasi
    
    # Isi surat
    lines = cover_letter.split('\n')
    in_body = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if 'Dengan hormat,' in line or 'Dear Sir/Madam,' in line:
            in_body = True
            doc.add_paragraph(line)
        elif 'Hormat saya,' in line or 'Sincerely,' in line:
            doc.add_paragraph()
            doc.add_paragraph(line)
        elif in_body and line and not line.startswith(nama.split()[0]):
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0.25)
        else:
            doc.add_paragraph(line)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

def create_pdf_file(cover_letter, perusahaan, posisi, nama):
    """Create PDF file dengan formatting profesional"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='Justify',
        parent=styles['Normal'],
        alignment=0,
        firstLineIndent=24
    ))
    
    # Build the PDF content
    story = []
    
    # Header
    story.append(Paragraph(nama, styles["Normal"]))
    story.append(Paragraph(DEFAULT_DATA['alamat'], styles["Normal"]))
    story.append(Paragraph(f"{DEFAULT_DATA['telepon']} | {DEFAULT_DATA['email']} | {DEFAULT_DATA['linkedin']}", styles["Normal"]))
    story.append(Paragraph(f"Portfolio: {DEFAULT_DATA['portofolio']}", styles["Normal"]))
    story.append(Spacer(1, 12))
    
    # Date
    story.append(Paragraph(datetime.date.today().strftime("%d %B %Y"), styles["Normal"]))
    story.append(Spacer(1, 12))
    
    # Recipient
    story.append(Paragraph("Kepada Yth:", styles["Normal"]))
    story.append(Paragraph("Manajer Rekrutmen", styles["Normal"]))
    story.append(Paragraph(perusahaan, styles["Normal"]))
    story.append(Spacer(1, 12))
    
    # Subject
    story.append(Paragraph(f"<b>Subjek: Lamaran Posisi {posisi}</b>", styles["Normal"]))
    story.append(Spacer(1, 12))
    
    # Body
    lines = cover_letter.split('\n')
    in_body = False
    
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
            
        if 'Dengan hormat,' in line or 'Dear Sir/Madam,' in line:
            in_body = True
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 6))
        elif 'Hormat saya,' in line or 'Sincerely,' in line:
            story.append(Spacer(1, 12))
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 12))
        elif in_body and line and not line.startswith(nama.split()[0]):
            story.append(Paragraph(line, styles["Justify"]))
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 6))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer

def get_download_link(file_buffer, filename, file_format):
    """Generate download link untuk file"""
    b64 = base64.b64encode(file_buffer.getvalue()).decode()
    href = f'<a href="data:application/{file_format};base64,{b64}" download="{filename}" style="background-color: #4CAF50; color: white; padding: 10px 20px; text-align: center; text-decoration: none; display: inline-block; border-radius: 4px; font-weight: bold;">📥 Download {file_format.upper()}</a>'
    return href

def analyze_match_score_professional(cover_letter, job_description, perusahaan, posisi, nama, skills, experiences):
    """Analisis match score dengan prompt profesional super detail"""
    
    prompt = PROFESSIONAL_PROMPTS["analyze_match"].format(
        posisi=posisi,
        perusahaan=perusahaan,
        nama=nama,
        applicant_skills=skills,
        applicant_experiences=experiences,
        cl_length=len(cover_letter),
        jd_length=len(job_description),
        job_description_analysis=extract_key_elements(job_description),
        cover_letter_content=cover_letter[:2000]
    )
    
    analysis = generate_with_gemini(prompt)
    return analysis

def enhance_cover_letter_professional(cover_letter, job_description, perusahaan, posisi, nama, match_analysis, skills, experiences):
    """Tingkatkan cover letter dengan prompt profesional super detail"""
    
    prompt = PROFESSIONAL_PROMPTS["enhance_cover"].format(
        posisi=posisi,
        perusahaan=perusahaan,
        nama=nama,
        applicant_skills=skills,
        applicant_experiences=experiences,
        original_cover_letter=cover_letter,
        job_description_analysis=extract_key_elements(job_description),
        match_analysis=match_analysis[:1000]
    )
    
    enhanced = generate_with_gemini(prompt)
    return enhanced

def improve_section_professional(section_content, section_type, job_description, perusahaan, posisi, skills, experiences):
    """Tingkatkan section tertentu dengan prompt profesional"""
    
    prompt = PROFESSIONAL_PROMPTS["improve_section"].format(
        posisi=posisi,
        perusahaan=perusahaan,
        section_type=section_type,
        applicant_skills=skills,
        applicant_experiences=experiences,
        content_length=len(section_content),
        original_content=section_content,
        job_keywords=extract_keywords(job_description)
    )
    
    improved = generate_with_gemini(prompt)
    return improved

def extract_key_elements(job_description):
    """Extract key elements from job description for analysis"""
    if not job_description:
        return "No job description provided"
    
    elements = []
    
    requirements = re.findall(r'(requirement|qualification|skill|ability|experience)[^.]*\.', job_description, re.IGNORECASE)
    if requirements:
        elements.append("KEY REQUIREMENTS: " + "; ".join(requirements[:3]))
    
    responsibilities = re.findall(r'(responsibilit|duties|will)[^.]*\.', job_description, re.IGNORECASE)
    if responsibilities:
        elements.append("KEY RESPONSIBILITIES: " + "; ".join(responsibilities[:3]))
    
    technical_terms = re.findall(r'\b(python|sql|java|javascript|machine learning|data analysis|web development|database)\b', job_description, re.IGNORECASE)
    if technical_terms:
        elements.append("TECHNICAL SKILLS MENTIONED: " + ", ".join(set(technical_terms)))
    
    return "\n".join(elements) if elements else "Standard professional requirements expected"

def extract_keywords(job_description):
    """Extract important keywords from job description"""
    if not job_description:
        return "No specific keywords identified"
    
    important_words = re.findall(r'\b([A-Z][a-z]+ [A-Z][a-z]+|[A-Z][a-z]+ing\b|\\b[A-Z][a-z]+ed\\b)', job_description)
    technical_words = re.findall(r'\b(python|sql|java|javascript|react|node|machine learning|data analysis|database|api|cloud|aws|azure)\b', job_description, re.IGNORECASE)
    
    keywords = list(set(important_words + technical_words))
    return ", ".join(keywords[:10]) if keywords else "General professional skills"

def parse_match_score(analysis_text):
    """Parse score dari hasil analisis profesional"""
    if not analysis_text:
        return 0
    
    score_patterns = [
        r'OVERALL MATCH SCORE:\s*(\d+)/100',
        r'SCORE:\s*(\d+)%',
        r'MATCH SCORE:\s*(\d+)',
        r'(\d+)/100'
    ]
    
    for pattern in score_patterns:
        match = re.search(pattern, analysis_text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    
    percent_pattern = r'(\d+)%'
    matches = re.findall(percent_pattern, analysis_text)
    if matches:
        return int(matches[0])
    
    return 0

def translate_text_safe(text, dest_language='en'):
    """Terjemahkan teks dengan error handling yang aman"""
    try:
        if not text or not text.strip():
            return text
            
        if TRANSLATOR_AVAILABLE:
            translated = GoogleTranslator(source='id', target=dest_language).translate(text)
            return translated
        else:
            return text
            
    except Exception as e:
        st.warning(f"Translasi gagal, menggunakan teks asli: {e}")
        return text

def generate_cover_letter_indonesia(data, perusahaan, posisi, konten, skills, experiences):
    """Generate surat lamaran versi Indonesia"""
    
    pembuka_formatted = konten['pembuka'].format(
        posisi=posisi, perusahaan=perusahaan, 
        jurusan=data['jurusan'], universitas=data['universitas']
    )
    
    alasan_formatted = konten['alasan_perusahaan'].format(perusahaan=perusahaan)
    penutup_formatted = konten['motivasi_penutup'].format(perusahaan=perusahaan)
    
    # Gabungkan pengalaman untuk konten yang lebih kaya
    pengalaman_text = " ".join(experiences['pengalaman_akademik'][:2] + experiences['pengalaman_organisasi'][:1])
    
    cover_letter = f"""
{data['nama']}
{data['alamat']}
{data['telepon']} | {data['email']} | {data['linkedin']}
Portfolio: {data['portofolio']}

{datetime.date.today().strftime("%d %B %Y")}

Kepada Yth:
Manajer Rekrutmen
{perusahaan}

Subjek: Lamaran Posisi {posisi}

Dengan hormat,

{pembuka_formatted}

{konten['pencapaian']} {pengalaman_text}

Keahlian teknis yang saya kuasai meliputi: {', '.join(skills['keahlian_teknis'][:5])}. 
Didukung dengan soft skills: {', '.join(skills['keahlian_lunak'][:4])} yang siap saya terapkan 
dalam menyelesaikan tantangan profesional.

{alasan_formatted}

{penutup_formatted}

Saya telah melampirkan CV dan dokumen pendukung untuk pertimbangan Bapak/Ibu. 
Saya sangat berharap dapat diberi kesempatan wawancara untuk membahas lebih lanjut 
bagaimana saya dapat berkontribusi bagi kemajuan {perusahaan}.

Hormat saya,

{data['nama']}
"""
    return cover_letter

def generate_cover_letter_english(data, perusahaan, posisi, konten, skills, experiences):
    """Generate surat lamaran versi Inggris"""
    
    pembuka_formatted = konten['pembuka'].format(
        posisi=posisi, perusahaan=perusahaan,
        jurusan=data['jurusan'], universitas=data['universitas']
    )
    
    alasan_formatted = konten['alasan_perusahaan'].format(perusahaan=perusahaan)
    penutup_formatted = konten['motivasi_penutup'].format(perusahaan=perusahaan)
    
    pembuka_en = translate_text_safe(pembuka_formatted)
    pencapaian_en = translate_text_safe(konten['pencapaian'])
    alasan_en = translate_text_safe(alasan_formatted)
    penutup_en = translate_text_safe(penutup_formatted)
    
    # Translate skills
    keahlian_teknis_en = [translate_text_safe(skill) for skill in skills['keahlian_teknis'][:5]]
    keahlian_lunak_en = [translate_text_safe(skill) for skill in skills['keahlian_lunak'][:4]]
    
    cover_letter = f"""
{data['nama']}
{data['alamat']}
{data['telepon']} | {data['email']} | {data['linkedin']}
Portfolio: {data['portofolio']}

{datetime.date.today().strftime("%B %d, %Y")}

To:
Recruitment Manager
{perusahaan}

Subject: Application for {posisi} Position

Dear Sir/Madam,

{pembuka_en}

{pencapaian_en}

My technical expertise includes: {', '.join(keahlian_teknis_en)}.
Supported by soft skills: {', '.join(keahlian_lunak_en)} that I am ready to apply 
in solving professional challenges.

{alasan_en}

{penutup_en}

I have attached my CV and supporting documents for your consideration.
I sincerely hope to be given an interview opportunity to further discuss 
how I can contribute to the progress of {perusahaan}.

Sincerely,

{data['nama']}
"""
    return cover_letter

def main():
    st.set_page_config(page_title="AI Cover Letter Pro", page_icon="🚀", layout="wide")
    
    st.title("🚀 Professional Cover Letter Generator")
    st.markdown("**Dibuat khusus untuk Kevin Jonathan Sinaga**")
    
    if not TRANSLATOR_AVAILABLE:
        st.warning("Fitur translasi terbatas. Install: `pip install deep-translator`")
    
    ai_ready = False
    
    with st.sidebar:
        st.header("⚙️ Konfigurasi")
        use_ai = st.checkbox("Aktifkan AI Generator", value=False)
        
        if use_ai:
            try:
                GEMINI_API_KEY = "AIzaSyCLXDKjPmL8XeJdmza7wZOWsIK93SHoUcI"
                genai.configure(api_key=GEMINI_API_KEY)
                ai_ready = True
                st.success("Konfigurasi AI Berhasil!")
            except Exception as e:
                st.error(f"Error Konfigurasi Gemini: {str(e)}")

    # Data pribadi
    st.header("👤 Data Pribadi")
    col1, col2 = st.columns(2)
    
    with col1:
        nama = st.text_input("Nama Lengkap*", value=DEFAULT_DATA['nama'])
        telepon = st.text_input("Telepon*", value=DEFAULT_DATA['telepon'])
        email = st.text_input("Email*", value=DEFAULT_DATA['email'])
        linkedin = st.text_input("LinkedIn", value=DEFAULT_DATA['linkedin'])
    
    with col2:
        alamat = st.text_area("Alamat*", value=DEFAULT_DATA['alamat'])
        jurusan = st.text_input("Jurusan*", value=DEFAULT_DATA['jurusan'])
        universitas = st.text_input("Universitas*", value=DEFAULT_DATA['universitas'])
        portofolio = st.text_input("Portfolio", value=DEFAULT_DATA['portofolio'])
    
    # Input Keahlian dan Pengalaman
    st.header("🎯 Keahlian & Pengalaman")
    st.info("Edit keahlian dan pengalaman Anda sesuai dengan background sebenarnya")
    
    col_skills, col_exp = st.columns(2)
    
    with col_skills:
        st.subheader("💻 Keahlian Teknis")
        keahlian_teknis = st.text_area(
            "Keahlian Teknis (pisahkan dengan koma)",
            value=", ".join(DEFAULT_SKILLS['keahlian_teknis']),
            height=100,
            help="Contoh: Python, Data Analysis, Web Development, SQL, Machine Learning Basics"
        )
        
        st.subheader("🌟 Soft Skills")
        keahlian_lunak = st.text_area(
            "Soft Skills (pisahkan dengan koma)",
            value=", ".join(DEFAULT_SKILLS['keahlian_lunak']),
            height=100,
            help="Contoh: Problem Solving, Team Leadership, Communication, Fast Learner"
        )
    
    with col_exp:
        st.subheader("📚 Pengalaman Akademik")
        pengalaman_akademik = st.text_area(
            "Pengalaman Akademik",
            value="\n".join(DEFAULT_EXPERIENCES['pengalaman_akademik']),
            height=120,
            help="Pengalaman selama kuliah, proyek tugas akhir, asisten lab, dll."
        )
        
        st.subheader("👥 Pengalaman Organisasi")
        pengalaman_organisasi = st.text_area(
            "Pengalaman Organisasi & Kepemimpinan",
            value="\n".join(DEFAULT_EXPERIENCES['pengalaman_organisasi']),
            height=120,
            help="Pengalaman memimpin tim, organisasi kampus, volunteer, dll."
        )
    
    # Data lowongan
    st.header("🏢 Informasi Lowongan")
    col3, col4 = st.columns(2)
    with col3:
        perusahaan = st.text_input("Nama Perusahaan*", placeholder="PT. Teknologi Indonesia", key="perusahaan")
    with col4:
        posisi = st.text_input("Posisi yang Dilamar*", placeholder="Data Analyst Junior", key="posisi")
    
    # Job Description Analysis Section
    st.header("📋 Job Description Analysis")
    st.info("Salin dan tempel iklan lowongan/reskripsi posisi dari perusahaan untuk analisis kesesuaian")
    
    job_description = st.text_area(
        "Job Description / Iklan Lowongan*",
        placeholder="Contoh: Kami mencari Data Analyst dengan kualifikasi:\n- Pengalaman Python dan SQL\n- Kemampuan analisis data\n- Komunikasi yang baik\n- Team player...",
        height=150,
        key="job_description"
    )
    
    # Konten surat
    st.header("📝 Konten Profesional")
    st.success("Konten sudah disiapkan otomatis. Anda bisa edit manual atau gunakan AI untuk improve!")
    
    # Initialize session state
    session_defaults = {
        'pembuka': DEFAULT_CONTENT['pembuka'],
        'pencapaian': DEFAULT_CONTENT['pencapaian'],
        'alasan_perusahaan': DEFAULT_CONTENT['alasan_perusahaan'],
        'motivasi_penutup': DEFAULT_CONTENT['motivasi_penutup'],
        'cover_id': None,
        'cover_en': None,
        'match_analysis': None,
        'match_score': 0,
        'enhanced_version': None
    }
    
    for key, value in session_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value
    
    col5, col6 = st.columns(2)
    
    with col5:
        st.subheader("🎯 Pembuka Profesional")
        if ai_ready and st.button("✨ Improve dengan AI", key="btn_pembuka"):
            with st.spinner("Memperbaiki pembuka..."):
                skills_data = {
                    'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                    'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
                }
                exp_data = {
                    'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                    'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
                }
                improved = improve_section_professional(
                    st.session_state.pembuka, "Opening Paragraph", job_description, perusahaan, posisi, skills_data, exp_data
                )
                if improved:
                    st.session_state.pembuka = improved
                    st.rerun()
        
        st.text_area("Kalimat Pembuka*", height=100, key="pembuka")
        
        st.subheader("💡 Pencapaian & Pengalaman")
        if ai_ready and st.button("✨ Improve dengan AI", key="btn_pencapaian"):
            with st.spinner("Memperbaiki pencapaian..."):
                skills_data = {
                    'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                    'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
                }
                exp_data = {
                    'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                    'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
                }
                improved = improve_section_professional(
                    st.session_state.pencapaian, "Achievements Section", job_description, perusahaan, posisi, skills_data, exp_data
                )
                if improved:
                    st.session_state.pencapaian = improved
                    st.rerun()
        
        st.text_area("Pencapaian Relevan*", height=120, key="pencapaian")
    
    with col6:
        st.subheader("🏢 Alasan Memilih Perusahaan")
        if ai_ready and st.button("✨ Improve dengan AI", key="btn_alasan"):
            with st.spinner("Memperbaiki alasan..."):
                skills_data = {
                    'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                    'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
                }
                exp_data = {
                    'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                    'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
                }
                improved = improve_section_professional(
                    st.session_state.alasan_perusahaan, "Company Alignment", job_description, perusahaan, posisi, skills_data, exp_data
                )
                if improved:
                    st.session_state.alasan_perusahaan = improved
                    st.rerun()
        
        st.text_area("Alasan Bergabung*", height=100, key="alasan_perusahaan")
        
        st.subheader("🎯 Motivasi & Penutup")
        if ai_ready and st.button("✨ Improve dengan AI", key="btn_penutup"):
            with st.spinner("Memperbaiki penutup..."):
                skills_data = {
                    'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                    'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
                }
                exp_data = {
                    'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                    'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
                }
                improved = improve_section_professional(
                    st.session_state.motivasi_penutup, "Closing Paragraph", job_description, perusahaan, posisi, skills_data, exp_data
                )
                if improved:
                    st.session_state.motivasi_penutup = improved
                    st.rerun()
        
        st.text_area("Penutup Powerful*", height=100, key="motivasi_penutup")
    
    # Tombol generate cover letter
    st.markdown("---")
    col_generate, col_analyze, col_enhance = st.columns([2, 1, 1])
    
    with col_generate:
        if st.button("🎯 GENERATE COVER LETTER", type="primary", use_container_width=True):
            data = {
                'nama': nama, 'telepon': telepon, 'email': email, 'alamat': alamat,
                'linkedin': linkedin, 'jurusan': jurusan, 'universitas': universitas,
                'portofolio': portofolio
            }
            
            # Process skills and experiences
            skills_data = {
                'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
            }
            exp_data = {
                'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
            }
            
            required_fields = [nama, telepon, email, alamat, jurusan, universitas, perusahaan, posisi]
            required_content = [
                st.session_state.pembuka, st.session_state.pencapaian, 
                st.session_state.alasan_perusahaan, st.session_state.motivasi_penutup
            ]
            
            if not all(required_fields) or not all(required_content):
                st.error("❌ Harap lengkapi semua field yang wajib diisi!")
            else:
                konten_final = {
                    "pembuka": st.session_state.pembuka,
                    "pencapaian": st.session_state.pencapaian,
                    "alasan_perusahaan": st.session_state.alasan_perusahaan,
                    "motivasi_penutup": st.session_state.motivasi_penutup
                }

                with st.spinner("Sedang generate cover letter profesional..."):
                    st.session_state.cover_id = generate_cover_letter_indonesia(data, perusahaan, posisi, konten_final, skills_data, exp_data)
                    st.session_state.cover_en = generate_cover_letter_english(data, perusahaan, posisi, konten_final, skills_data, exp_data)
                    st.session_state.enhanced_version = None
                
                st.success("✅ Cover Letter profesional berhasil dibuat!")
    
    with col_analyze:
        if st.session_state.cover_id and job_description and ai_ready:
            if st.button("📊 ANALYZE MATCH SCORE", use_container_width=True):
                with st.spinner("Menganalisis kesesuaian profesional..."):
                    skills_data = {
                        'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                        'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
                    }
                    exp_data = {
                        'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                        'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
                    }
                    st.session_state.match_analysis = analyze_match_score_professional(
                        st.session_state.cover_id, job_description, perusahaan, posisi, nama, skills_data, exp_data
                    )
                    st.session_state.match_score = parse_match_score(st.session_state.match_analysis)
                    st.rerun()
    
    with col_enhance:
        if st.session_state.cover_id and st.session_state.match_analysis and ai_ready:
            if st.button("🚀 ENHANCE COVER LETTER", use_container_width=True):
                with st.spinner("Meningkatkan kualitas profesional..."):
                    skills_data = {
                        'keahlian_teknis': [s.strip() for s in keahlian_teknis.split(',')],
                        'keahlian_lunak': [s.strip() for s in keahlian_lunak.split(',')]
                    }
                    exp_data = {
                        'pengalaman_akademik': [e.strip() for e in pengalaman_akademik.split('\n') if e.strip()],
                        'pengalaman_organisasi': [e.strip() for e in pengalaman_organisasi.split('\n') if e.strip()]
                    }
                    enhanced = enhance_cover_letter_professional(
                        st.session_state.cover_id, job_description, perusahaan, posisi, nama, st.session_state.match_analysis, skills_data, exp_data
                    )
                    if enhanced:
                        st.session_state.enhanced_version = enhanced
                        st.success("✅ Cover Letter berhasil ditingkatkan secara profesional!")
                        st.rerun()
    
    # Tampilkan hasil jika ada
    if st.session_state.cover_id:
        st.markdown("---")
        st.header("📄 Hasil Cover Letter")
        
        # Tampilkan Match Score jika sudah dianalisis
        if st.session_state.match_analysis:
            st.subheader(f"🎯 Professional Match Score: {st.session_state.match_score}/100")
            
            score_color = "🟢" if st.session_state.match_score >= 80 else "🟡" if st.session_state.match_score >= 60 else "🔴"
            st.progress(st.session_state.match_score / 100, text=f"{score_color} Professional Alignment: {st.session_state.match_score}%")
            
            with st.expander("📋 Detailed Professional Analysis"):
                st.markdown(st.session_state.match_analysis)
        
        # Tabs untuk berbagai versi
        tab_names = ["🇮🇩 Versi Indonesia"]
        if st.session_state.enhanced_version:
            tab_names.insert(0, "🚀 Enhanced Version")
        if st.session_state.cover_en:
            tab_names.append("🇺🇸 English Version")
        
        tabs = st.tabs(tab_names)
        
        current_tab = 0
        
        # Enhanced Version Tab
        if st.session_state.enhanced_version:
            with tabs[current_tab]:
                st.subheader("🚀 Enhanced Professional Version")
                st.info("Versi yang telah dioptimalkan secara profesional berdasarkan analisis match score")
                st.text_area("Enhanced Cover Letter", st.session_state.enhanced_version, height=500, key="enhanced_result")
                
                # Download buttons untuk enhanced version
                col_docx1, col_pdf1 = st.columns(2)
                with col_docx1:
                    docx_buffer = create_docx_file(st.session_state.enhanced_version, perusahaan, posisi, nama)
                    st.markdown(get_download_link(docx_buffer, f"Enhanced_Cover_Letter_{perusahaan.replace(' ', '_')}.docx", "vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)
                
                with col_pdf1:
                    pdf_buffer = create_pdf_file(st.session_state.enhanced_version, perusahaan, posisi, nama)
                    st.markdown(get_download_link(pdf_buffer, f"Enhanced_Cover_Letter_{perusahaan.replace(' ', '_')}.pdf", "pdf"), unsafe_allow_html=True)
                
                current_tab += 1
        
        # Indonesian Version Tab
        with tabs[current_tab]:
            st.text_area("Cover Letter Indonesia", st.session_state.cover_id, height=500, key="id_result")
            
            # Download buttons untuk Indonesian version
            col_docx2, col_pdf2 = st.columns(2)
            with col_docx2:
                docx_buffer = create_docx_file(st.session_state.cover_id, perusahaan, posisi, nama)
                st.markdown(get_download_link(docx_buffer, f"Cover_Letter_ID_{perusahaan.replace(' ', '_')}.docx", "vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)
            
            with col_pdf2:
                pdf_buffer = create_pdf_file(st.session_state.cover_id, perusahaan, posisi, nama)
                st.markdown(get_download_link(pdf_buffer, f"Cover_Letter_ID_{perusahaan.replace(' ', '_')}.pdf", "pdf"), unsafe_allow_html=True)
            
            current_tab += 1
        
        # English Version Tab
        if st.session_state.cover_en:
            with tabs[current_tab]:
                st.text_area("Cover Letter English", st.session_state.cover_en, height=500, key="en_result")
                
                # Download buttons untuk English version
                col_docx3, col_pdf3 = st.columns(2)
                with col_docx3:
                    docx_buffer = create_docx_file(st.session_state.cover_en, perusahaan, posisi, nama)
                    st.markdown(get_download_link(docx_buffer, f"Cover_Letter_EN_{perusahaan.replace(' ', '_')}.docx", "vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)
                
                with col_pdf3:
                    pdf_buffer = create_pdf_file(st.session_state.cover_en, perusahaan, posisi, nama)
                    st.markdown(get_download_link(pdf_buffer, f"Cover_Letter_EN_{perusahaan.replace(' ', '_')}.pdf", "pdf"), unsafe_allow_html=True)

if __name__ == "__main__":
    main()